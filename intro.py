import cv2 as cv
import numpy as np 
import matplotlib.pyplot as plt 
from PIL import Image 
import pytesseract
from tesserocr import PyTessBaseAPI, PSM, OEM
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from skimage import measure, feature
from skimage.segmentation import clear_border
import argparse
import pickle
import os
import warnings
warnings.filterwarnings("ignore")
def bold_text(image):
    img = cv.erode(image, np.ones((7,7)), 2)
    img = cv.dilate(img, np.ones((7,7)))
    hist = cv.reduce(img, 0, cv.REDUCE_AVG).ravel()
    size = len(hist)
    start = np.where(hist > 0)[0][0]
    if hist[start] > 30:
        th = 12.5
    elif hist[start] > 20:
        th = 10.5
    elif hist[start] > 10:
        th = 8.5
    elif hist[start] > 5:
        th = 6.5
    else:
        th = 4.2
    for i in range(start, size - 25):
        if np.mean(hist[i : i + 25]) < th:
            break
    return image[:, : i + 5]
def display(image, cmap = None):
    plt.figure(figsize = (15, 15))
    plt.imshow(image, cmap = cmap)
def preprocessing(image):
    gray = cv.cvtColor(image, cv.COLOR_BGR2GRAY)
    #gray = cv.GaussianBlur(gray, (7,7), 0)
    gray = cv.bilateralFilter(gray, 4, 70, 70)
    thresh = cv.threshold(gray, 0, 255, cv.THRESH_BINARY_INV | cv.THRESH_OTSU)[1]
    return thresh
def _skew(image):
    pts = cv.findNonZero(image)
    if pts is None or len(pts) == 0:
        return None, None
    center, (w, h), angle = cv.minAreaRect(pts)
    if angle < - 45:
        angle += 90
    M = cv.getRotationMatrix2D(center, angle, 1.)
    image = cv.warpAffine(image, M, image.shape[:2][::-1])
    return image, M 
def split_page(image):
    w = image.shape[1]
    return (image[:, : w // 2], image[:, w // 2 : ])
def remove_border(image):
    image[: 60] = 0
    image[image.shape[0] - 60 :] = 0,
    image[:, : 150] = 0
    image[:, image.shape[1] - 60 : ] = 0
    return image
def remove_num_page(img):
    hist = cv.reduce(img, 1, cv.REDUCE_AVG).ravel()
    size = img.shape[0] - 1
    for i in range(size, 0, -1):
        if hist[i] > 20:
            break
    return i
def remove_character(img):
    hist = cv.reduce(img, 0, cv.REDUCE_AVG).ravel()
    idx = np.where(hist > 0)[0]
    start = idx[0]
    idx = idx[-1]
    thresh = 4
    i = start
    for i in range(idx, start, -1):
        if hist[i - 1] > thresh and hist[i] < thresh:
            break
    if i < img.shape[1] // 2:
        return img
    return img[..., : i - 1]
def skew(pages, p_pages):
    n_pages, n_p_pages = [], []
    for page, p_page in zip(pages, p_pages):
        pp, mp = _skew(p_page)
        n_p_pages.append(pp)
        n_pages.append(cv.warpAffine(page, mp, page.shape[:2][::-1]))
    return n_pages, n_p_pages
def split_column(image):
  hist = cv.reduce(image, 0, cv.REDUCE_AVG).ravel()
  idx = np.where(hist > 0)[0][0]
  for i in range(idx, hist.shape[0] - 20):
    if hist[i : i + 20].mean() < 10:
      break
  tp = (image[:, idx - 10 : i + 15],)
  sub_col = image[:, i + 15 :]
  num_non_zero = cv.countNonZero(sub_col)
  if num_non_zero < 50:
    return tp
  return tp + (sub_col,)
def split_row(img):
    dilate_img = cv.dilate(img, np.ones((25, 25)), 20)
    hist = cv.reduce(dilate_img, 1, cv.REDUCE_AVG).ravel()
    th = 10
    h, w = img.shape
    lowers, uppers = [], []
    for i in range(h - 1):
        if hist[i] < th and hist[i + 1] >= th:
            lowers.append(i)
        elif hist[i] >= th and hist[i + 1] < th:
            uppers.append(i)
    lines = []
    for l, u in zip(lowers, uppers):
        line = img[l - 6 : u + 6]
        pts = cv.countNonZero(line)
        if pts is None or pts < 100:
          continue
        lines.append(line)
    return lines
def split_line(img):
  hist = cv.reduce(img, 1, cv.REDUCE_AVG).ravel()
  th = 12
  h, w = img.shape
  lowers, uppers = [], []
  for i in range(h - 1):
    if hist[i] < th and hist[i + 1] >= th:
      lowers.append(i)
    elif hist[i] >=th and hist[i + 1] < th:
      uppers.append(i)
  lines = []
  for l, u in zip(lowers, uppers):
    line = img[l - 9: u + 9]
    h, w = line.shape[:2]
    if h < 30:
      continue
    lines.append(line)
  return lines
def split_word(image):
    padding = 2
    thresh = 7
    kernel = np.ones((5, 5))
    img = cv.dilate(image, kernel)
    hist = cv.reduce(img, 0, cv.REDUCE_AVG).ravel()
    id = np.where(hist > 0)[0][0]
    sep = [id]
    h, w = image.shape
    for i in range(w - 1):
        if hist[i] > thresh and hist[i + 1] < thresh:
            sep.append(i + 1)
    img_ls = []
    if hist[-1] > thresh:
        sep.append(w)
    sep.append(0)
    size = len(sep)
    for i in range(size - 1):
        roi = image[padding:, sep[i] - padding : sep[i + 1] + padding]
        r_h, r_w = roi.shape
        if r_w < 15:
            continue
        img_ls.append(roi)
    return img_ls
def split_character(word):
    h, w = word.shape
    # labels = measure.label(word, neighbors = 8, background = 0)
    labels = measure.label(word, neighbors = 8, connectivity = 2)
    c = word
    min_x = w
    for label in np.unique(labels):
        if label == 0:
            continue
        mask = np.zeros(word.shape, dtype = np.uint8)
        mask[labels == label] = 255
        pts = cv.findNonZero(mask)
        _, (_w, _h), _ = cv.minAreaRect(pts)
        if _h < 15:
            continue
        x, y = np.where(mask > 0)
        x, y = sorted(x), sorted(y)
        x_start, y_start = max(0, x[0] - 4), max(0, y[0] - 4)
        x_end, y_end = min(h, x[-1] + 4), min(w, y[-1] + 4)
        if y_end - y_start + 1 > 35:
            continue
        img = word[x_start : x_end, y_start : y_end]
        if img.shape[1] > c.shape[1] or x_start < min_x:
            min_x = x_start
            c = img
    return c
def generate_type_word(d):
    dt = {}
    for i in d:
        dt[i + ')'] = d[i] + ")"
        dt[i + ').'] = d[i] + ")."
        dt['(' + i] = "(" + d[i] 
        dt['(' + i + ')'] = '(' + d[i] + ')'
        dt['(' + i + ').'] = '(' + d[i] + ').'
        dt[i] = d[i]
    return dt
def load_model():
    with open('classification.pickle', 'rb') as f:
        clf = pickle.loads(f.read())
    with open('scaler.pickle', 'rb') as f:
        sc = pickle.loads(f.read())
    return clf, sc

def is_italic(clf, sc, image):
    image = cv.resize(image, (40, 40))
    vec = feature.hog(image, orientations = 9, pixels_per_cell = (4,4), cells_per_block = (2,2), transform_sqrt = True, block_norm = 'L1').reshape(1, -1)
    vec = sc.transform(vec)
    return clf.predict(vec)
def typeofword(d, text):
    dt = {'cv.': 'cũng viết', 'cn.': 'cũng nói', 'd.': 'danh từ', 't.': 'tính từ', 'tr.': 'trợ từ', 'đ.': 'đại từ', 'x.':'xem', 'đg.': 'động từ','p.':'phụ từ'}
    # for c in dt:
    #     idx = text.find(c)
    #     if idx != -1 and (text[idx - 1] == ' ' or text[idx - 1] == '"') and text[idx + len(c)] != ')':
    #         words = text.split(c)
    #         return (words[0].split(), dt[c], words[1].split())
    # return (text.split(),)
    words = text.split()
    size = len(words)
    for i in range(size):
        if words[i] in dt:
            return ([d.get(i,i) for i in words[:i]], dt[words[i]], [d.get(i,i) for i in words[i + 1:]])
    return ([d.get(i,i) for i in words], )
        
def isEndOfDefine(img):
    x, y = np.where(img)
    minY, maxY = min(y), max(y)
    return maxY - minY <= 687

def parseargument():
    ap = argparse.ArgumentParser()
    ap.add_argument('-i', '--image', required = True, type = str, help = 'path to image')
    ap.add_argument('-t', '--intro', required = True, type = str, help = 'Are both 2 of pages about introduction')
    args = vars(ap.parse_args())
    return args

def main(args):
    d = {'b.': 'bóng (nghĩa bóng)', 'c.': 'cảm từ', 'cd.': 'ca dao', 'chm.': 'chuyên môn', 'id.': 'ít dùng', 'k.': 'kết từ', 'kc.': 'kiểu cách',
        'kng.': 'khẩu ngữ', 'ng.': 'nghĩa', 'ph.': 'phương ngữ', 'thgt.': 'thông tục', 'tng.': 'tục ngữ', 'trtr.': 'trang trọng', 'vch.': 'văn chương',
        'cv.': 'cũng viết', 'cn.': 'cũng nói', 'd.': 'danh từ', 't.': 'tính từ', 'tr.': 'trợ từ', 'đ.': 'đại từ', 'x.':'xem', 
        'đg.': 'động từ','p.':'phụ từ', 'đp.': 'động từ', 'đợ.': 'động từ', 'đẹ.': 'động từ'}
        
    d = generate_type_word(d)

    clf, sc = load_model()
    img_path = args['image']
    is_all_intro = args['intro'][0] == 'T'
    img = cv.imread(img_path)
    thresh = preprocessing(img)
    pages = split_page(img)
    p_pages = [preprocessing(page) for page in pages]
    p_pages = [remove_border(page) for page in p_pages]
    pages, p_pages = skew(pages, p_pages)
    document = Document()
    footer = document.sections[0].footer
    font = document.styles['Normal'].font
    font.name = 'Time New Roman'
    font.size = Pt(9)
    first = True
    num = False
    if is_all_intro:
        count = 2
    else:
        count = 1
    with PyTessBaseAPI(lang = 'vie', psm = PSM.SINGLE_LINE, oem = OEM.LSTM_ONLY) as api:
        n_page = 0
        for ppage in p_pages[: count]:
            p = document.add_paragraph('')
            rois = split_row(ppage)[:-1]
            for roi in rois:
                if roi.shape[0] < 90:
                    rows = [roi]
                else:
                    rows = split_line(roi)
                for row in rows:
                    api.SetImage(Image.fromarray(cv.bitwise_not(row)))
                    text = api.GetUTF8Text().strip()
                    if len(text) == 0:
                        continue
                    print(text)
                    tus = split_word(row)
                    l_i = []
                    for tu in tus:
                        c = split_character(tu)
                        #t_c = pytesseract.image_to_string(cv.bitwise_not(tu), lang = 'vie', config = '--oem 1 --psm 7')
                        api.SetImage(Image.fromarray(cv.bitwise_not(tu)))
                        t_c = api.GetUTF8Text()
                        l_i.extend([is_italic(clf, sc, c)[0]] * len(t_c.split()))
                    words = text.split()
                    for item in words:
                        try:
                            p.add_run('{} '.format(int(item))).bold = True
                        except:
                            if len(l_i) and l_i[0] == 1:
                                p.add_run('{} '.format(item)).italic = True
                            else:
                                p.add_run('{} '.format(item))
                        if len(l_i):
                            l_i.pop(0)
                    p.add_run('\n')
            if n_page < count - 1:
                p.add_run().add_break(WD_BREAK.PAGE)
            n_page += 1
        #p = document.add_paragraph('')
        if not is_all_intro:
            p.add_run().add_break(WD_BREAK.PAGE)
            roi = split_row(p_pages[1])[1]
            cols = split_column(roi)
            for col in cols:
                i = remove_num_page(col)
                col = col[: i + 12]
                col = _skew(col)[0]
                if col is None:
                    continue
                rows = split_line(col)
                if len(rows) == 0:
                    continue
                ls = []
                for row in rows:
                    l_i = []
                    api.SetImage(Image.fromarray(cv.bitwise_not(row)))
                    text = api.GetUTF8Text().strip()
                    if len(text) == 0:
                        continue
                    tus = split_word(row)
                    for tu in tus:
                        c = split_character(tu)
                        api.SetImage(Image.fromarray(cv.bitwise_not(tu)))
                        t_c = api.GetUTF8Text()
                        l_i.extend([is_italic(clf, sc, c)[0]] * len(t_c.split()))
                    ls_words = typeofword(d, text)
                    if num or first:
                        if len(ls):
                            for l in ls:
                                if l[1] == 1:
                                    p.add_run('{} '.format(l[0])).bold = True
                                elif l[1] == 2:
                                    p.add_run('{} '.format(l[0])).italic = True
                                else:
                                    p.add_run('{} '.format(l[0]))
                            p.add_run('\n')
                        ls = []
                        if len(ls_words) > 1:
                            for b_word in ls_words[0]:
                                ls.append((b_word, 1))
                                if len(l_i):
                                    l_i.pop(0)
                            ls.append((ls_words[1], 2))
                            if len(l_i):
                                l_i.pop(0)
                            ls_words = (ls_words[2], )
                        else:
                            size = len(ls_words[0])
                            i = 0
                            nfound = False
                            for i in range(size):
                                item = ls_words[0][i]
                                if (item[0] != '[' and item[0] != '(' and not item[0].isupper()) or (len(item) == 1 and item[0].isalpha()):
                                    ls.append((item, 1))
                                else:
                                    nfound = True
                                    break
                                if len(l_i):
                                    l_i.pop(0)
                            if nfound:
                                ls_words = (ls_words[0][i:], )
                            else:
                                ls_words = ([], )
                    else:
                        if len(ls_words) > 1:
                            temp = [ls_words[1]]
                            temp.extend(ls_words[2])
                            ls_words[0].extend(temp)
                    first = False
                    for item in ls_words[0]:
                        try:
                            if int(item) < 11:
                                ls.append((item, 1))
                            elif len(l_i) and l_i[0] == 1:
                                ls.append((item, 2))
                            else:
                                ls.append((item, 0))
                        except:
                            if len(l_i) and l_i[0] == 1:
                                ls.append((item, 2))
                            else:
                                ls.append((item, 0))
                        if len(l_i):
                            l_i.pop(0)    
                    num = isEndOfDefine(row)
                if len(ls):
                    for l in ls:
                        if l[1] == 1:
                            p.add_run('{} '.format(l[0])).bold = True
                        elif l[1] == 2:
                            p.add_run('{} '.format(l[0])).italic = True
                        else:
                            p.add_run('{} '.format(l[0]))
                    if len(ls):
                        p.add_run('\n')
    document.save('output/{}.docx'.format(args['image'].split(os.path.sep)[-1].split('.')[0]))

if __name__ == "__main__":
    main(parseargument())

